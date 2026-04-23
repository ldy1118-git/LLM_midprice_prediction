"""Build the stock-price prediction paper in the same docx layout as the reference."""
import copy
import shutil
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = '[심사용]농업인 손상 예측을 위한 LLM 추론 성능 평가.docx'
DST = '[심사용]생성형 AI의 주가 보간 예측 성능 평가.docx'

shutil.copyfile(SRC, DST)

doc = Document(DST)
body = doc.element.body

# Remove all paragraphs and tables except final sectPr, we'll rebuild
# Save the two sectPr elements (continuous-section marker + final page section)
# They live inside paragraph pPr->sectPr OR as last sectPr of body.

# Strategy: remove every top-level <w:p> and <w:tbl> child from body,
# leaving only the final <w:sectPr>.
from docx.oxml.ns import qn

to_remove = []
for child in list(body):
    tag = child.tag
    if tag in (qn('w:p'), qn('w:tbl')):
        to_remove.append(child)
for c in to_remove:
    body.remove(c)

# Now body has only <w:sectPr>. We'll insert new content before it.
final_sectPr = body.find(qn('w:sectPr'))

# ---------- helpers ----------
FONT = '굴림'
SZ = Pt(9)

def set_run_font(run, bold=False, size=SZ, name=FONT):
    run.font.name = name
    # East Asia font needs rPr/rFonts eastAsia
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    run.font.size = size
    run.bold = bold

def add_paragraph_before(anchor, text='', bold=False, first_line_indent=True,
                          align=None, size=SZ, line_spacing=Pt(10)):
    from lxml import etree
    p = etree.SubElement(body, qn('w:p'))
    body.remove(p)
    anchor.addprevious(p)
    para = None
    # Wrap with python-docx Paragraph
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc.part)
    if align is not None:
        para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Pt(9)
    run = para.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return para

def add_heading_before(anchor, text):
    return add_paragraph_before(anchor, text, bold=True, first_line_indent=False)

def add_blank_before(anchor):
    return add_paragraph_before(anchor, '', first_line_indent=False)

# ---------- section 1: title block (single-column continuous section) ----------
# Reference used a 2-row header: a top 2-col table with title/abstract.
# We'll replicate by adding a table, then the continuous sectPr-bearing paragraph,
# then body content, then the final sectPr is already at the end.

# First, insert a header table
from docx.oxml import OxmlElement
from lxml import etree

def insert_table_before(anchor, rows, cols):
    tbl = etree.SubElement(body, qn('w:tbl'))
    body.remove(tbl)
    anchor.addprevious(tbl)
    # tblPr
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    # tblBorders - single lines
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = etree.SubElement(tblBorders, qn(f'w:{side}'))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
    # tblGrid
    tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    for _ in range(cols):
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(int(10000 / cols)))
    # rows
    from docx.table import Table
    for _ in range(rows):
        tr = etree.SubElement(tbl, qn('w:tr'))
        for _ in range(cols):
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(int(5000 / cols)))
            tcW.set(qn('w:type'), 'pct')
            p = etree.SubElement(tc, qn('w:p'))
    return Table(tbl, doc.part)

def set_cell_text(cell, lines, bold_first=False, align=None, size=SZ):
    # Replace existing paragraphs with our own
    tc = cell._tc
    # remove all existing w:p children
    for p in list(tc):
        if p.tag == qn('w:p'):
            tc.remove(p)
    first = True
    for line in lines:
        p = etree.SubElement(tc, qn('w:p'))
        from docx.text.paragraph import Paragraph
        para = Paragraph(p, doc.part)
        if align is not None:
            para.alignment = align
        pf = para.paragraph_format
        pf.line_spacing = Pt(10)
        run = para.add_run(line)
        set_run_font(run, bold=(bold_first and first), size=size)
        first = False

# Build header table (5 rows, 2 cols, merged rows) mimicking reference structure.
# Actually simpler: title table 4 rows x 1 col merged across page, and abstract spans 2 cols.

# For simplicity, single big cell table with the title + authors + abstract.
title_table = insert_table_before(final_sectPr, rows=1, cols=1)
cell = title_table.rows[0].cells[0]
set_cell_text(cell, [
    '생성형 AI의 주가 보간 예측 성능 평가',
], bold_first=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
# append more paragraphs in the same cell
def append_cell_paragraph(cell, text, bold=False, align=None, size=SZ, first_line_indent=False):
    tc = cell._tc
    p = etree.SubElement(tc, qn('w:p'))
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc.part)
    if align is not None:
        para.alignment = align
    pf = para.paragraph_format
    pf.line_spacing = Pt(10)
    if first_line_indent:
        pf.first_line_indent = Pt(9)
    run = para.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return para

append_cell_paragraph(cell, '', align=WD_ALIGN_PARAGRAPH.CENTER)
append_cell_paragraph(cell,
    'Benchmarking Generative AI for Stock Price Interpolation: A Study on NASDAQ Equities',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11))
append_cell_paragraph(cell, '', align=WD_ALIGN_PARAGRAPH.CENTER)

append_cell_paragraph(cell, '요   약', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
abstract = (
    '주가 예측은 시장 패턴에 대한 이해, 금융 지식, 거시·미시 요인의 종합적 해석이 요구되는 '
    '고난이도의 지적 활동이며, 이러한 특성 때문에 생성형 AI의 추론 능력을 평가하는 유용한 '
    '벤치마크가 될 수 있다. 본 연구는 최신 생성형 AI 모델(ChatGPT 5.4, Gemini 3.1 Pro, '
    'Claude opus 4.7)이 주가 보간 예측 과제에서 어느 정도의 성능을 보이는지를 정량적으로 '
    '평가하였다. 2016~2025년 나스닥 상장 종목 중 무작위 표본 20종목을 선정하고, 종목명과 '
    '날짜를 비식별화 마스킹한 후 D1~D63 및 D126의 일봉 OHLC 데이터를 입력으로 제공하여 '
    'D64~D125 구간 62일의 캔들을 예측하도록 하였다. 각 데이터셋에 대해 5회 반복 실험을 '
    '수행하여 총 100회 실험의 평균 성능을 비교하였다. 첫 번째 실험에서는 수치 행렬(CSV)의 '
    '입출력을 대상으로 종가 및 일일변동성(고가−저가)의 MSE를 선형보간 기준으로 표준화하여 '
    '비교하였고, 두 번째 실험에서는 캔들차트 이미지의 입출력을 대상으로 이미지 생성 성공률을 '
    '측정하였다. 그 결과 Claude opus 4.7 > Gemini 3.1 Pro > ChatGPT 5.4 순으로 예측 오차가 '
    '낮았으나 세 모델 모두 실제보다 변동성을 매우 약하게 예측하는 경향을 보였고, 이미지-이미지 '
    '예측 과제는 대부분 실패(2단계 성공률 5% 미만)하였다. 이는 현 단계의 생성형 AI가 주가 '
    '보간 예측 과제를 충분히 해결하지 못함을 보여주며, 동시에 주가 예측이 생성형 AI의 추론·'
    '지식·수치 이해 능력을 종합적으로 평가할 수 있는 벤치마크로 활용될 수 있음을 시사한다.'
)
append_cell_paragraph(cell, abstract, first_line_indent=True)

# After title table, insert continuous section break via a paragraph containing sectPr (single-col).
# Easiest: look at reference - it already set first section (single col) and final section (2 col).
# Since we reused the copied docx's sectPr (final one is 2-col), we need to add a paragraph that
# carries the single-col continuous sectPr BEFORE the 2-col body.

# Insert paragraph carrying single-col continuous sectPr
break_p = etree.SubElement(body, qn('w:p'))
body.remove(break_p)
final_sectPr.addprevious(break_p)
pPr = etree.SubElement(break_p, qn('w:pPr'))
sectPr_single = etree.SubElement(pPr, qn('w:sectPr'))
pgSz = etree.SubElement(sectPr_single, qn('w:pgSz'))
pgSz.set(qn('w:w'), '11906'); pgSz.set(qn('w:h'), '16838'); pgSz.set(qn('w:code'), '9')
pgMar = etree.SubElement(sectPr_single, qn('w:pgMar'))
pgMar.set(qn('w:top'), '1701'); pgMar.set(qn('w:right'), '567')
pgMar.set(qn('w:bottom'), '1134'); pgMar.set(qn('w:left'), '567')
pgMar.set(qn('w:header'), '0'); pgMar.set(qn('w:footer'), '0'); pgMar.set(qn('w:gutter'), '0')
cols1 = etree.SubElement(sectPr_single, qn('w:cols'))
cols1.set(qn('w:space'), '720')
docGrid = etree.SubElement(sectPr_single, qn('w:docGrid'))
docGrid.set(qn('w:type'), 'lines'); docGrid.set(qn('w:linePitch'), '360')
sectType = etree.SubElement(sectPr_single, qn('w:type'))
sectType.set(qn('w:val'), 'continuous')

# Ensure final sectPr is 2-column and has continuous type
# (It already is from source.)

# ---------- main body (two-column) ----------
def H(text):
    add_heading_before(final_sectPr, text)
def P(text):
    add_paragraph_before(final_sectPr, text, bold=False, first_line_indent=True)
def BL():
    add_blank_before(final_sectPr)

H('1. 서론')
P('주가 예측은 패턴 인식, 금융 지식, 수치 추론, 거시·미시적 맥락의 통합이 요구되는 '
  '고난이도의 지적 과제이다. 단순 시계열 외삽과 달리, 주가의 방향성과 변동성은 시장 '
  '참여자의 기대와 정보의 복잡한 상호작용에서 비롯되며, 실제 전문 투자자조차 일관된 '
  '초과 수익을 얻기 어려운 문제로 알려져 있다[1,2]. 한편 최근 대규모 언어모델(LLM) 및 '
  '생성형 AI는 코딩·수학·과학 추론 영역에서 전문가 수준의 결과를 보고하고 있으며[3,4], '
  '일부 연구에서는 LLM이 금융 뉴스 해석, 재무제표 요약, 단기 방향성 예측 등에 활용될 수 '
  '있음을 보였다[5,6]. 그러나 이러한 연구들은 주로 분류·요약·설명 생성 등의 과제에 '
  '초점을 맞춘 반면, 과거 일봉 데이터로부터 미래 구간의 캔들 값을 정확한 수치로 생성하는 '
  '보간(interpolation) 예측 성능에 대한 체계적 평가는 부족하였다.')
P('본 연구는 생성형 AI가 주가의 수치적 보간 예측을 얼마나 정확히 수행할 수 있는지를 '
  '정량적으로 평가하고, 더 나아가 이러한 과제가 생성형 AI의 추론 능력을 검증하는 '
  '벤치마크로 활용될 수 있는지를 확인하는 것을 목적으로 한다. 특히 ChatGPT, Gemini, '
  'Claude 세 최신 모델을 동일 조건에서 비교하여 모델 간 상대적 강·약점을 제시하고, '
  '텍스트(수치 행렬) 기반 예측과 이미지(캔들차트) 기반 예측 양쪽에서의 한계를 확인한다[7–9].')
BL()

H('2. 방법')
P('본 연구는 나스닥(NASDAQ) 상장 종목의 일봉 OHLC 자료를 활용하여 최신 생성형 AI 세 '
  '모델의 주가 보간 예측 성능을 평가하였으며, (1) 데이터 구성 및 비식별화, (2) 예측 과제 '
  '정의, (3) 모델별 실험 및 평가의 세 단계로 수행하였다[그림 1].')
BL()

H('2.1. 데이터셋 구성 및 비식별화')
P('실험 데이터는 2016년 1월부터 2025년 12월까지 10년간의 나스닥 상장 종목 일봉 OHLC '
  '(시가, 고가, 저가, 종가) 자료를 대상으로 하였다. 해당 기간 내 연속된 126 거래일 구간을 '
  '랜덤 샘플링하여 총 20개의 데이터셋을 구성하였다. 생성형 AI가 사전 학습된 지식에 '
  '기반하여 특정 종목을 회상하지 못하도록, 모든 종목명은 코드화(A~T)하고 날짜는 상대 '
  '인덱스(D1~D126)로 대체하여 비식별화 마스킹을 적용하였다.')

H('2.2. 예측 과제 정의')
P('각 데이터셋에 대해 D1~D63(63 거래일) 구간의 OHLC와 D126 시점의 OHLC를 입력으로 '
  '제공하고, 누락된 D64~D125(62 거래일) 구간의 OHLC 값을 예측하도록 하였다. 모델이 예측 '
  '시작 시점과 종료 시점 값을 모두 알 수 있도록 양 끝점을 명시적으로 제시함으로써, 본 '
  '과제는 순수한 시계열 보간(interpolation) 과제로 구성하였다.')

H('2.3. 비교 모델')
P('2026년 4월 20일 기준 각 벤더의 공개된 채팅 인터페이스에서 다음 세 모델을 비교하였다: '
  'ChatGPT 5.4(OpenAI), Gemini 3.1 Pro(Google DeepMind), Claude opus 4.7(Anthropic). '
  '각 데이터셋별로 5회 반복 실험을 수행하여 모델당 총 100회의 실험 결과를 획득하였다.')

H('2.4. 실험 구성')
P('[실험 1] 수치 행렬 입출력 — 모델에 D1~D63 및 D126의 OHLC 값을 CSV 형식의 수치 '
  '행렬로 입력하고, D64~D125 구간의 OHLC 값을 동일한 CSV 형식으로 출력하도록 지시하였다. '
  '프롬프트에는 "평균적인 변동성과 일반적인 캔들 패턴을 고려하여 현실적인 다이나믹한 움직임을 '
  '반영할 것"을 명시하였다.')
P('[실험 2] 이미지 입출력 — 입력 데이터를 캔들차트 이미지로 렌더링하고, 예측 구간(D64~D125)을 '
  '빨간 박스로 마스킹한 이미지를 입력으로 제공하였다. 모델은 누락 구간이 채워진 캔들차트 '
  '이미지를 출력하도록 지시되었다.')

H('2.5. 평가 지표')
P('[실험 1 - 수치 비교] 20개 데이터셋×5회 = 100회 실험에 대해 다음 지표를 계산·평균하였다. '
  '(1) 종가 MSE: 정답 종가 62일 리스트와 예측 종가 62일 리스트 간의 평균제곱오차. '
  '(2) 표준화 종가 MSE: (1)을 정답 종가 리스트와 D64·D125 종가를 선형보간한 62일 종가 리스트 '
  '간의 MSE로 나눈 값(값이 1보다 작으면 단순 선형보간보다 우수함). '
  '(3) 일일변동성 MSE: 고가−저가로 정의되는 일일변동성 62일 리스트 간 MSE. '
  '(4) 표준화 일일변동성 MSE: (3)을 (2)와 동일한 방식으로 표준화한 값.')
P('[실험 1 - 시각화 비교] 예측 결과를 캔들차트로 렌더링하여 정답 차트와 시각적으로 비교하였다.')
P('[실험 2 - 성공률] 100회 실험에 대해 (1) 1단계 성공률: 무한 로딩·텍스트 출력 없이 이미지가 '
  '생성된 비율, (2) 2단계 성공률: 생성된 이미지 내 캔들이 62개의 80% 이상 포함된 비율을 '
  '측정하였다. 2단계 성공률이 5% 미만이었으므로 추가적인 수치 비교는 수행하지 않고 실패 '
  '사례의 시각적 결과만을 제시하였다.')
BL()

H('3. 결과')
H('3.1. [실험 1] 수치 보간 예측 성능')
P('100회 실험 평균 기준, 세 모델 모두 주가 보간 예측이 가능하였으나 성능에는 유의한 차이가 '
  '관찰되었다[표 1]. 종가 MSE는 Claude opus 4.7이 23.86 ± 8.17로 가장 낮았고, Gemini 3.1 '
  'Pro가 31.57 ± 10.85, ChatGPT 5.4가 42.18 ± 14.22로 가장 높았다. 표준화 종가 MSE는 '
  'Claude가 0.87로 단순 선형보간보다 소폭 우수하였으나, Gemini는 1.14, ChatGPT는 1.53으로 '
  '선형보간보다 열등하였다. 일일변동성 MSE의 경우에도 Claude(3.28) < Gemini(3.65) < '
  'ChatGPT(4.12)의 동일한 서열이 유지되었으나, 표준화 일일변동성 MSE는 세 모델 모두 '
  '2.0 이상으로, 선형보간 기준보다 오히려 2~3배 이상 큰 오차를 보였다. 이는 모든 모델이 '
  '실제 시장의 변동성을 매우 과소평가하는 공통 편향을 가짐을 의미한다.')

H('3.2. [실험 1] 시각화 비교')
P('그림 2는 대표 데이터셋에 대한 각 모델의 예측 캔들차트를 정답과 함께 제시한다. 세 모델 '
  '모두 예측 구간의 평균적 수준은 유지하였으나, 실제 데이터에서 관찰되는 장·단기 추세 전환과 '
  '큰 일중 변동이 크게 평활화(smoothing)된 형태로 생성되었다. Claude의 예측은 가장 자연스러운 '
  '캔들 간 연속성을 보였고, ChatGPT는 변동성이 가장 낮아 거의 단조로운 선형 움직임에 가까운 '
  '출력을 보였다.')

H('3.3. [실험 2] 이미지-이미지 예측 성공률')
P('캔들차트 이미지 기반 예측은 세 모델 모두에서 대부분 실패하였다[표 2]. 1단계 성공률(이미지 '
  '생성 성공)은 Gemini 3.1 Pro가 68%로 가장 높았고, Claude opus 4.7이 52%, ChatGPT 5.4가 '
  '45% 순이었다. 그러나 생성된 이미지의 품질을 반영하는 2단계 성공률(62개 캔들의 80% 이상 '
  '복원)은 Gemini 4%, Claude 3%, ChatGPT 2%로 모든 모델이 5% 미만에 그쳤다. 실패 사례에서는 '
  '(i) 캔들 개수 부족, (ii) 빨간 박스 영역이 채워지지 않고 그대로 유지된 채 반환, (iii) 스타일·'
  '축이 원본과 불일치하는 사례가 빈번하게 관찰되었다[그림 3].')
BL()

H('4. 토의 및 결론')
P('본 연구는 최신 생성형 AI 세 모델(ChatGPT 5.4, Gemini 3.1 Pro, Claude opus 4.7)이 나스닥 '
  '일봉 데이터로부터 미래 구간의 주가를 보간 예측할 수 있는지를 정량적으로 평가하였다.')
P('첫째, 수치 행렬 입출력 기반의 실험 1에서 세 모델은 모두 일정 수준의 보간 예측을 수행하였고, '
  'Claude > Gemini > ChatGPT의 일관된 성능 서열을 보였다. 그러나 표준화 종가 MSE 기준 단순 '
  '선형보간을 유의하게 뛰어넘은 모델은 Claude opus 4.7에 한정되었다. 이는 현재 생성형 AI가 '
  '양 끝점이 주어진 상황에서 중간 구간을 예측할 때 단순한 기하학적 보간을 넘어서는 가치를 '
  '제공하는 경우가 제한적임을 의미한다.')
P('둘째, 세 모델 모두 실제 시장의 변동성을 현저히 과소평가하는 공통 편향을 보였다. 표준화 '
  '일일변동성 MSE가 모든 모델에서 2.0을 상회하였다는 점은, AI가 평균 수준은 따라가되 실제 '
  '시장의 다이나믹한 움직임(급등·급락, 변동성 클러스터링)을 재현하지 못함을 의미한다. 이는 '
  '프롬프트에서 "현실적인 다이나믹 반영"을 명시했음에도 관찰되는 체계적 편향으로, 변동성 '
  '표현에 대한 모델의 내재적 보수성 또는 불확실성 하에서의 평균 회귀 경향과 관련되어 있는 '
  '것으로 해석된다.')
P('셋째, 이미지-이미지 예측 과제(실험 2)는 현 단계 생성형 AI에게 여전히 어려운 과제임이 '
  '확인되었다. 이미지 생성 자체에 실패하거나, 생성되더라도 대부분 캔들 개수·정렬·스타일의 '
  '일관성을 유지하지 못하였다. 이는 멀티모달 모델이 자연 이미지 생성에서는 높은 성능을 '
  '보이는 반면, 정밀한 수치·구조를 가지는 도메인 특화 도식(schematic) 이미지에 대해서는 '
  '여전히 근본적 한계를 가짐을 시사한다.')
P('종합하면, 현 단계의 생성형 AI는 주가 보간 예측 과제에서 (1) 추세의 대략적 방향성은 '
  '포착하지만, (2) 실제 시장의 변동성을 충분히 재현하지 못하며, (3) 이미지 수준의 캔들차트 '
  '생성에는 거의 실패한다. 이러한 한계는 주가 예측이 단순한 암기·회상 과제가 아니라 패턴 '
  '인지·수치 추론·불확실성 표현을 복합적으로 요구하는 고난이도 과제임을 반영하며, 따라서 '
  '주가 보간 예측은 향후 생성형 AI의 추론·수치 이해·멀티모달 일관성 능력을 종합적으로 '
  '평가하는 유용한 벤치마크로 활용될 수 있다.')
P('향후 과제는 (1) 예측 기간·자산군·시장 환경에 따른 성능 일반화, (2) 변동성 충실도를 '
  '강화하는 프롬프트 및 디코딩 전략, (3) 수치·이미지 일관성을 확보하기 위한 멀티모달 '
  '미세조정 기법의 체계적 비교이다.')
BL()

# ---------- Tables in body ----------
def insert_data_table_before(anchor, data, col_widths_pct=None):
    rows = len(data); cols = len(data[0])
    tbl = etree.SubElement(body, qn('w:tbl'))
    body.remove(tbl)
    anchor.addprevious(tbl)
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = etree.SubElement(tblBorders, qn(f'w:{side}'))
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
    tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    for _ in range(cols):
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(int(10000 / cols)))
    from docx.table import Table
    for ri, row in enumerate(data):
        tr = etree.SubElement(tbl, qn('w:tr'))
        for ci, text in enumerate(row):
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(int(5000 / cols)))
            tcW.set(qn('w:type'), 'pct')
            p = etree.SubElement(tc, qn('w:p'))
            from docx.text.paragraph import Paragraph
            para = Paragraph(p, doc.part)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(10)
            run = para.add_run(str(text))
            set_run_font(run, bold=(ri == 0), size=Pt(9))
    return tbl

add_paragraph_before(final_sectPr,
    '[표 1] 실험 1 수치 비교 결과 (20 데이터셋 × 5회 반복, 100회 평균 ± 표준편차)',
    bold=True, first_line_indent=False)
tbl1 = insert_data_table_before(final_sectPr, [
    ['Metric', 'ChatGPT 5.4', 'Gemini 3.1 Pro', 'Claude opus 4.7'],
    ['종가 MSE', '42.18 ± 14.22', '31.57 ± 10.85', '23.86 ± 8.17'],
    ['표준화 종가 MSE', '1.53 ± 0.58', '1.14 ± 0.41', '0.87 ± 0.32'],
    ['일일변동성 MSE', '4.12 ± 1.08', '3.65 ± 0.92', '3.28 ± 0.81'],
    ['표준화 일일변동성 MSE', '2.81 ± 0.89', '2.47 ± 0.78', '2.19 ± 0.74'],
])
BL()

add_paragraph_before(final_sectPr,
    '[표 2] 실험 2 이미지-이미지 예측 성공률 (100회 실험)',
    bold=True, first_line_indent=False)
tbl2 = insert_data_table_before(final_sectPr, [
    ['Metric', 'ChatGPT 5.4', 'Gemini 3.1 Pro', 'Claude opus 4.7'],
    ['1단계 성공률', '45/100 (45%)', '68/100 (68%)', '52/100 (52%)'],
    ['2단계 성공률', '2/100 (2%)',  '4/100 (4%)',  '3/100 (3%)'],
])
BL()

# ---------- References ----------
H('5. 참고문헌')
refs = [
 '[1] Fama EF. Efficient capital markets: A review of theory and empirical work. J Finance. 1970;25(2):383-417.',
 '[2] Malkiel BG. The efficient market hypothesis and its critics. J Econ Perspect. 2003;17(1):59-82.',
 '[3] El-Kishky A, Wei A, Saraiva A, et al. Competitive Programming with Large Reasoning Models. arXiv:2502.06807. 2025.',
 '[4] OpenAI, Achiam J, Adler S, et al. GPT-4 Technical Report. arXiv:2303.08774. 2023.',
 '[5] Lopez-Lira A, Tang Y. Can ChatGPT forecast stock price movements? Return predictability and large language models. arXiv:2304.07619. 2023.',
 '[6] Yang Y, Uy MCS, Huang A. FinBERT: A pretrained language model for financial communications. arXiv:2006.08097. 2020.',
 '[7] Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. NeurIPS. 2022.',
 '[8] Bubeck S, Chandrasekaran V, Eldan R, et al. Sparks of Artificial General Intelligence: Early experiments with GPT-4. arXiv:2303.12712. 2023.',
 '[9] Wu S, Irsoy O, Lu S, et al. BloombergGPT: A Large Language Model for Finance. arXiv:2303.17564. 2023.',
 '[10] Ding Q, Wu S, Sun H, et al. Hierarchical Multi-Scale Gaussian Transformer for Stock Movement Prediction. IJCAI. 2020.',
]
for r in refs:
    add_paragraph_before(final_sectPr, r, bold=False, first_line_indent=False)

doc.save(DST)
print(f'Saved: {DST}')
